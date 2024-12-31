from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def ajustar_espacamento(paragrafo, espaco_antes=0, espaco_depois=0):
    """
    Ajusta o espaçamento antes e depois de um parágrafo.

    Args:
        paragrafo (docx.text.paragraph.Paragraph): O parágrafo a ser ajustado.
        espaco_antes (int): Espaço antes do parágrafo (em pontos).
        espaco_depois (int): Espaço depois do parágrafo (em pontos).
    """
    p = paragrafo._element
    pPr = p.get_or_add_pPr()
    spacing = pPr.get_or_add_spacing()
    spacing.set(qn('w:before'), str(espaco_antes))
    spacing.set(qn('w:after'), str(espaco_depois))

def adicionar_secao(document, titulo, conteudo=None, espaco_antes_titulo=200, espaco_depois_titulo=50, espaco_entre_conteudo=100):
    """
    Adiciona uma seção ao documento com espaçamento ajustável.

    Args:
        document (docx.Document): O documento Word.
        titulo (str): Título da seção.
        conteudo (str ou list): Conteúdo da seção (string ou lista de strings).
        espaco_antes_titulo (int): Espaço antes do título.
        espaco_depois_titulo (int): Espaço depois do título.
        espaco_entre_conteudo (int): Espaço entre itens do conteúdo.
    """
    titulo_paragrafo = document.add_heading(titulo, level=1)
    ajustar_espacamento(titulo_paragrafo, espaco_antes=espaco_antes_titulo, espaco_depois=espaco_depois_titulo)

    if conteudo:
        if isinstance(conteudo, list):
            for item in conteudo:
                paragrafo = document.add_paragraph(item, style='List Bullet')
                ajustar_espacamento(paragrafo, espaco_antes=0, espaco_depois=espaco_entre_conteudo)
        else:
            paragrafo = document.add_paragraph(conteudo)
            ajustar_espacamento(paragrafo, espaco_antes=0, espaco_depois=espaco_entre_conteudo)

def criar_documento_curso(curso_titulo, curso_descricao, curso_professor, topicos, calendario, nome_arquivo=None):
    """
    Cria um documento Word com as informações do curso.

    Args:
        curso_titulo (str): Título do curso.
        curso_descricao (str): Descrição do curso.
        curso_professor (str): Nome do professor.
        topicos (list): Lista de tópicos da ementa.
        calendario (list): Lista de tuplas (data, tópico) para o calendário.
        nome_arquivo (str ou None): Nome do arquivo para salvar o documento (opcional).
    """
    document = Document()

    # Adiciona título do curso
    titulo = document.add_heading(curso_titulo, level=0)
    ajustar_espacamento(titulo, espaco_antes=0, espaco_depois=150)

    # Adiciona descrição do curso
    descricao = document.add_paragraph(curso_descricao)
    ajustar_espacamento(descricao, espaco_antes=0, espaco_depois=100)

    # Adiciona seção do professor
    adicionar_secao(document, 'Professor', curso_professor, espaco_antes_titulo=200, espaco_depois_titulo=50)

    # Adiciona seção de ementa
    adicionar_secao(document, 'Ementa', topicos, espaco_antes_titulo=200, espaco_depois_titulo=50, espaco_entre_conteudo=50)

    # Adiciona seção do calendário
    adicionar_secao(document, 'Calendário', espaco_antes_titulo=200, espaco_depois_titulo=50)
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Configura espaçamento na tabela
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Data'
    hdr_cells[1].text = 'Tópico'

    for data, topico in calendario:
        row_cells = table.add_row().cells
        row_cells[0].text = data
        row_cells[1].text = topico

    # Ajusta espaçamento nas linhas da tabela
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '0')  # Reduz o espaço depois das linhas da tabela
        trPr.append(spacing)

    # Define o nome do arquivo se não for fornecido
    if not nome_arquivo:
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        nome_arquivo = f"documento_{timestamp}.docx"

    # Salvando o documento
    document.save(nome_arquivo)
    print(f"Documento salvo como {nome_arquivo}")

# Exemplo de uso
if __name__ == '__main__':
    curso_titulo = "Curso de Python Avançado"
    curso_descricao = "Aprenda conceitos avançados de Python para desenvolvimento profissional."
    curso_professor = "Antônio Abrantes"
    topicos = ["Programação Funcional", "Programação Assíncrona", "Testes Automatizados", "Otimização de Código", "Integração Contínua"]
    calendario = [
        ("01/01/2024", "Programação Funcional"),
        ("08/01/2024", "Programação Assíncrona"),
        ("15/01/2024", "Testes Automatizados"),
        ("22/01/2024", "Otimização de Código"),
        ("29/01/2024", "Integração Contínua"),
    ]
    criar_documento_curso(curso_titulo, curso_descricao, curso_professor, topicos, calendario)
